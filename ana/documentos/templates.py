"""Templates e estilos para geração de peças processuais com python-docx.

Aplica formatação compatível com as normas processuais brasileiras:
- Fonte: Arial 12pt (corpo) / 14pt (títulos)
- Margens: 3cm superior/esquerda, 2cm inferior/direita (padrão ABNT/forense)
- Espaçamento: 1,5 entre linhas
- Parágrafos com recuo de 2,5cm na primeira linha
"""

from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


def criar_documento_juridico() -> Document:
    """Cria um Document com estilos jurídicos aplicados."""
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Estilo Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    _set_spacing(normal, espaco_antes=0, espaco_depois=0, line_rule="multiple", line_val="276")

    # Estilo Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    _set_spacing(h1, espaco_antes=240, espaco_depois=120, line_rule="multiple", line_val="276")

    return doc


def _set_spacing(
    style,
    espaco_antes: int,
    espaco_depois: int,
    line_rule: str,
    line_val: str,
) -> None:
    """Define espaçamento de parágrafo no estilo."""
    pPr = style.element.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(espaco_antes))
    spacing.set(qn("w:after"), str(espaco_depois))
    spacing.set(qn("w:lineRule"), line_rule)
    spacing.set(qn("w:line"), line_val)
    pPr.append(spacing)


def add_cabecalho_enderecamento(doc: Document, enderecamento: str) -> None:
    """Adiciona parágrafo de endereçamento centralizado (ex: 'AO JUÍZO...')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(enderecamento.upper())
    run.bold = True
    run.font.size = Pt(12)


def add_qualificacao(doc: Document, texto: str) -> None:
    """Adiciona bloco de qualificação das partes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(2.5)
    p.add_run(texto)


def add_secao(doc: Document, numero: str, titulo: str, corpo: str) -> None:
    """Adiciona uma seção numerada (ex: 'I — DOS FATOS')."""
    doc.add_paragraph()
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(f"{numero} — {titulo.upper()}")

    for paragrafo in corpo.strip().split("\n\n"):
        paragrafo = paragrafo.strip()
        if not paragrafo:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(2.5)
        p.add_run(paragrafo)


def add_pedidos(doc: Document, pedidos: str) -> None:
    """Adiciona seção de pedidos com fórmula de encerramento."""
    add_secao(doc, "III", "DOS PEDIDOS", pedidos)

    doc.add_paragraph()
    p_fecho = doc.add_paragraph()
    p_fecho.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_fecho.paragraph_format.first_line_indent = Cm(2.5)
    p_fecho.add_run(
        "Nestes termos, pede deferimento."
    )


def add_local_data_assinatura(doc: Document, cidade_uf: str) -> None:
    """Adiciona local, data e espaço para assinatura."""
    hoje = date.today()
    data_fmt = f"{hoje.day} de {_mes_por_extenso(hoje.month)} de {hoje.year}"
    cidade = cidade_uf.split("/")[0].strip() if "/" in cidade_uf else cidade_uf

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{cidade}, {data_fmt}.")

    # Linha de assinatura
    doc.add_paragraph()
    doc.add_paragraph()
    p_linha = doc.add_paragraph()
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_linha.add_run("_" * 50)

    p_advogado = doc.add_paragraph()
    p_advogado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_advogado.add_run("Advogado(a)")


def add_numero_processo(doc: Document, numero: str, tipo_acao: str) -> None:
    """Insere bloco com número e tipo do processo."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PROCESSO Nº {numero}\n{tipo_acao.upper()}")
    run.bold = True


def _mes_por_extenso(mes: int) -> str:
    meses = [
        "", "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
    ]
    return meses[mes] if 1 <= mes <= 12 else str(mes)
