"""Gerador de peças processuais para o sistema ANA.

Orquestra RAG + LLM + python-docx para gerar peças jurídicas completas:
    - Petição Inicial
    - Contestação
    - Recurso
    - Exportação de Transcrição

Fluxo:
    1. Carrega dados da sessão (processo, partes, área)
    2. RAG: busca jurisprudência e legislação relevante
    3. LLM: gera conteúdo de cada seção com base nos dados + RAG
    4. python-docx: monta o documento com formatação processual
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from loguru import logger

from ana.documentos.modelos import TipoPeca, label_peca
from ana.documentos.templates import (
    add_cabecalho_enderecamento,
    add_local_data_assinatura,
    add_numero_processo,
    add_pedidos,
    add_qualificacao,
    add_secao,
    criar_documento_juridico,
)
from ana.sessoes.modelos import Sessao


def gerar_documento(
    sessao_id: str,
    tipo_peca: TipoPeca,
    instrucoes: str = "",
) -> tuple[bytes, str]:
    """Gera uma peça processual como documento Word.

    Args:
        sessao_id: ID da sessão do processo.
        tipo_peca: Tipo de peça a gerar.
        instrucoes: Instruções adicionais para o LLM.

    Returns:
        Tupla (bytes_do_docx, nome_do_arquivo).

    Raises:
        ValueError: Se a sessão não existir.
        RuntimeError: Se o LLM falhar.
    """
    from ana.sessoes.repositorio import obter_sessao

    sessao = obter_sessao(sessao_id)
    if sessao is None:
        raise ValueError(f"Sessão '{sessao_id}' não encontrada.")

    logger.info(
        f"ALERJ: gerando {label_peca(tipo_peca)} para processo {sessao.numero_processo}"
    )

    if tipo_peca == TipoPeca.EXPORTAR_TRANSCRICAO:
        docx_bytes = _gerar_exportacao_transcricao(sessao)
    elif tipo_peca == TipoPeca.PETICAO_INICIAL:
        docx_bytes = _gerar_peticao_inicial(sessao, instrucoes)
    elif tipo_peca == TipoPeca.CONTESTACAO:
        docx_bytes = _gerar_contestacao(sessao, instrucoes)
    elif tipo_peca == TipoPeca.RECURSO:
        docx_bytes = _gerar_recurso(sessao, instrucoes)
    else:
        raise ValueError(f"Tipo de peça não suportado: {tipo_peca}")

    numero_curto = sessao.numero_processo.replace("/", "-").replace(".", "-")[:20]
    nome = f"{tipo_peca.value}_{numero_curto}.docx"
    logger.info(f"Documento gerado: {nome} ({len(docx_bytes)} bytes)")
    return docx_bytes, nome


# ---------------------------------------------------------------------------
# Petição Inicial
# ---------------------------------------------------------------------------


def _gerar_peticao_inicial(sessao: Sessao, instrucoes: str) -> bytes:
    contexto_rag = _buscar_contexto_rag(sessao, k=6)
    llm = _obter_llm()

    fatos = _gerar_secao_llm(
        llm,
        sessao,
        secao="DOS FATOS",
        instrucoes=instrucoes,
        contexto_rag="",  # fatos vêm dos dados do processo, não do RAG
    )
    direito = _gerar_secao_llm(
        llm,
        sessao,
        secao="DO DIREITO",
        instrucoes=instrucoes,
        contexto_rag=contexto_rag,
    )
    pedidos = _gerar_secao_llm(
        llm,
        sessao,
        secao="DOS PEDIDOS",
        instrucoes=instrucoes,
        contexto_rag="",
    )

    doc = criar_documento_juridico()

    endereco = _montar_enderecamento(sessao)
    add_cabecalho_enderecamento(doc, endereco)
    doc.add_paragraph()
    add_numero_processo(doc, sessao.numero_processo, sessao.tipo_acao)
    doc.add_paragraph()
    add_qualificacao(doc, _montar_qualificacao(sessao))
    add_secao(doc, "I", "DOS FATOS", fatos)
    add_secao(doc, "II", "DO DIREITO", direito)
    add_pedidos(doc, pedidos)
    add_local_data_assinatura(doc, sessao.cidade_uf)

    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Contestação
# ---------------------------------------------------------------------------


def _gerar_contestacao(sessao: Sessao, instrucoes: str) -> bytes:
    contexto_rag = _buscar_contexto_rag(sessao, k=6)
    llm = _obter_llm()

    preliminares = _gerar_secao_llm(
        llm, sessao, secao="PRELIMINARES", instrucoes=instrucoes, contexto_rag=""
    )
    merito = _gerar_secao_llm(
        llm, sessao, secao="DO MÉRITO", instrucoes=instrucoes, contexto_rag=contexto_rag
    )
    pedidos = _gerar_secao_llm(
        llm, sessao, secao="DOS PEDIDOS (CONTESTAÇÃO)", instrucoes=instrucoes, contexto_rag=""
    )

    doc = criar_documento_juridico()

    endereco = _montar_enderecamento(sessao)
    add_cabecalho_enderecamento(doc, endereco)
    doc.add_paragraph()
    add_numero_processo(doc, sessao.numero_processo, f"Contestação — {sessao.tipo_acao}")
    doc.add_paragraph()
    add_qualificacao(doc, _montar_qualificacao_contestacao(sessao))
    add_secao(doc, "I", "DAS PRELIMINARES", preliminares)
    add_secao(doc, "II", "DO MÉRITO", merito)
    add_pedidos(doc, pedidos)
    add_local_data_assinatura(doc, sessao.cidade_uf)

    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Recurso
# ---------------------------------------------------------------------------


def _gerar_recurso(sessao: Sessao, instrucoes: str) -> bytes:
    contexto_rag = _buscar_contexto_rag(sessao, k=6)
    llm = _obter_llm()

    cabimento = _gerar_secao_llm(
        llm, sessao, secao="DO CABIMENTO", instrucoes=instrucoes, contexto_rag=""
    )
    merito = _gerar_secao_llm(
        llm, sessao, secao="DO MÉRITO RECURSAL", instrucoes=instrucoes, contexto_rag=contexto_rag
    )
    pedido = _gerar_secao_llm(
        llm, sessao, secao="DO PEDIDO RECURSAL", instrucoes=instrucoes, contexto_rag=""
    )

    doc = criar_documento_juridico()

    add_cabecalho_enderecamento(doc, "EGRÉGIO TRIBUNAL DE JUSTIÇA")
    doc.add_paragraph()
    add_numero_processo(doc, sessao.numero_processo, f"Recurso — {sessao.tipo_acao}")
    doc.add_paragraph()
    add_qualificacao(doc, _montar_qualificacao(sessao))
    add_secao(doc, "I", "DO CABIMENTO", cabimento)
    add_secao(doc, "II", "DO MÉRITO", merito)
    add_pedidos(doc, pedido)
    add_local_data_assinatura(doc, sessao.cidade_uf)

    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Exportação de Transcrição
# ---------------------------------------------------------------------------


def _gerar_exportacao_transcricao(sessao: Sessao) -> bytes:
    """Gera DOCX formatado com a transcrição da sessão, se disponível."""
    from ana.sessoes.repositorio import listar_documentos

    docs = listar_documentos(sessao.id)
    transcricoes = [d for d in docs if "transcric" in d.nome.lower() or d.tipo == "transcricao"]

    doc = criar_documento_juridico()

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("TRANSCRIÇÃO DE AUDIÊNCIA")
    run.bold = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run(
        f"Processo: {sessao.numero_processo}\n"
        f"Tipo: {sessao.tipo_acao}\n"
        f"Exportado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
    )

    if transcricoes:
        doc.add_paragraph()
        doc.add_paragraph().add_run("— Conteúdo da Transcrição —").bold = True
        for t in transcricoes:
            doc.add_paragraph().add_run(f"Arquivo: {t.nome}")
    else:
        doc.add_paragraph()
        doc.add_paragraph().add_run(
            "Nenhuma transcrição indexada nesta sessão. "
            "Use o endpoint /transcricao para transcrever audiências."
        )

    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Helpers LLM
# ---------------------------------------------------------------------------


def _obter_llm():
    from ana.config import obter_configuracao
    from ana.config_modelos import obter_modelos
    from ana.providers.llm import OllamaLLMProvider

    config = obter_configuracao()
    modelo = obter_modelos().ativo.agentes.redator
    return OllamaLLMProvider(modelo=modelo, host=config.ollama_host)


def _gerar_secao_llm(
    llm,
    sessao: Sessao,
    secao: str,
    instrucoes: str,
    contexto_rag: str,
) -> str:
    """Usa o LLM redator para gerar o conteúdo de uma seção."""
    partes_txt = _formatar_partes(sessao.partes)

    prompt_partes = (
        f"Processo nº {sessao.numero_processo}\n"
        f"Tipo de ação: {sessao.tipo_acao}\n"
        f"Área jurídica: {sessao.area}\n"
        f"Vara: {sessao.vara or 'não informada'}\n"
        f"Foro: {sessao.cidade_uf or 'não informado'}\n"
        f"Partes: {partes_txt}\n"
    )

    rag_bloco = ""
    if contexto_rag:
        rag_bloco = (
            "\n\nFundamentação jurídica disponível (legislação e jurisprudência):\n"
            + contexto_rag
        )

    instrucoes_bloco = f"\n\nInstruções adicionais: {instrucoes.strip()}" if instrucoes.strip() else ""

    prompt = (
        "Você é um advogado brasileiro especialista em redação de peças processuais. "
        "Redija em português jurídico brasileiro formal e técnico. "
        "Seja objetivo, preciso e utilize linguagem processual adequada. "
        "Retorne apenas o texto da seção solicitada, sem títulos, sem explicações.\n\n"
        f"Dados do processo:\n{prompt_partes}"
        f"{rag_bloco}"
        f"{instrucoes_bloco}\n\n"
        f"Redija agora a seção '{secao}':"
    )

    try:
        resultado = llm.invocar(prompt, temperatura=0.3)
        return resultado.strip()
    except Exception as e:
        logger.warning(f"LLM falhou na seção '{secao}': {e}")
        return f"[Seção {secao} — aguardando revisão do advogado responsável.]"


# ---------------------------------------------------------------------------
# Helpers RAG
# ---------------------------------------------------------------------------


def _buscar_contexto_rag(sessao: Sessao, k: int = 6) -> str:
    """Busca legislação e jurisprudência relevante para o processo."""
    try:
        from ana.rag.retrieval import obter_pipeline_retrieval

        pipeline = obter_pipeline_retrieval()
        query = f"{sessao.tipo_acao} {sessao.area} {sessao.numero_processo}"
        chunks = pipeline.buscar(query=query, k=k)

        if not chunks:
            return ""

        trechos = []
        for i, chunk in enumerate(chunks, 1):
            fonte = getattr(chunk, "fonte", None) or (
                chunk.metadata.fonte if hasattr(chunk, "metadata") else "desconhecida"
            )
            texto = getattr(chunk, "texto", str(chunk))
            trechos.append(f"[{i}] {fonte}:\n{texto[:400]}")

        return "\n\n".join(trechos)
    except Exception as e:
        logger.warning(f"RAG indisponível para geração de documento: {e}")
        return ""


# ---------------------------------------------------------------------------
# Helpers de formatação
# ---------------------------------------------------------------------------


def _montar_enderecamento(sessao: Sessao) -> str:
    vara = sessao.vara or "VARA CÍVEL"
    cidade = sessao.cidade_uf or "COMARCA LOCAL"
    return f"EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(ÍZA) DE DIREITO\nDA {vara.upper()} DE {cidade.upper()}"


def _montar_qualificacao(sessao: Sessao) -> str:
    partes = sessao.partes
    autor = partes.get("autor") or partes.get("requerente") or "AUTOR (qualificação a completar)"
    reu = partes.get("reu") or partes.get("requerido") or partes.get("réu") or "RÉU (qualificação a completar)"

    return (
        f"{autor}, já devidamente qualificado nos autos do processo em epígrafe, "
        f"vem, por meio de seu(sua) advogado(a) subscritor(a), com fulcro no art. 319 do "
        f"Código de Processo Civil, propor a presente {sessao.tipo_acao} "
        f"em face de {reu}, pelos fatos e fundamentos jurídicos a seguir expostos."
    )


def _montar_qualificacao_contestacao(sessao: Sessao) -> str:
    partes = sessao.partes
    reu = partes.get("reu") or partes.get("requerido") or partes.get("réu") or "RÉU"
    autor = partes.get("autor") or partes.get("requerente") or "AUTOR"

    return (
        f"{reu}, já qualificado nos autos, vem, por meio de seu(sua) advogado(a), "
        f"tempestivamente, apresentar CONTESTAÇÃO à {sessao.tipo_acao} proposta por "
        f"{autor}, pelos fundamentos a seguir aduzidos."
    )


def _formatar_partes(partes: dict[str, Any]) -> str:
    if not partes:
        return "não informadas"
    return "; ".join(f"{k}: {v}" for k, v in partes.items())


def _doc_to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
